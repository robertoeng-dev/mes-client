"""
Gerador de documentação oficial do MES Client.
Execute: python gerar_documentacao.py
Gera: MES_Client_Documentacao_Oficial.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─── Cores ────────────────────────────────────────────────────────────────────
AZUL_ESCURO  = RGBColor(0x1A, 0x37, 0x6C)   # azul escuro institucional
AZUL_MEDIO   = RGBColor(0x27, 0x5D, 0xA6)   # azul médio
CINZA_TEXTO  = RGBColor(0x40, 0x40, 0x40)   # texto corpo
CINZA_CLARO  = RGBColor(0xF2, 0xF2, 0xF2)   # fundo de tabela (cabeçalho)
BRANCO       = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = AZUL_ESCURO
        run.font.size = Pt(18)
    elif level == 2:
        run.font.color.rgb = AZUL_MEDIO
        run.font.size = Pt(14)
    else:
        run.font.color.rgb = CINZA_TEXTO
        run.font.size = Pt(12)
    return p

def add_body(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = CINZA_TEXTO
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = CINZA_TEXTO
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Cabeçalho
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = BRANCO
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, "1A376C")

    # Linhas
    for r_idx, row in enumerate(rows):
        bg = "F2F2F2" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].runs[0].font.color.rgb = CINZA_TEXTO
            set_cell_bg(cell, bg)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)

    # fundo cinza claro via shading no parágrafo
    pPr   = p._p.get_or_add_pPr()
    shd   = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "EBEBEB")
    pPr.append(shd)
    return p

# ══════════════════════════════════════════════════════════════════════════════
doc = Document()

# Margens
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# Estilo padrão
style = doc.styles["Normal"]
style.font.name  = "Calibri"
style.font.size  = Pt(11)
style.font.color.rgb = CINZA_TEXTO

# ── CAPA ─────────────────────────────────────────────────────────────────────
doc.add_paragraph("\n\n\n")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MES CLIENT")
run.font.size  = Pt(36)
run.font.bold  = True
run.font.color.rgb = AZUL_ESCURO

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Parser & Monitor de Resultados de Teste")
run.font.size  = Pt(18)
run.font.color.rgb = AZUL_MEDIO

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Manual do Usuário e Documentação Técnica")
run.font.size  = Pt(14)
run.font.italic = True
run.font.color.rgb = CINZA_TEXTO

doc.add_paragraph("\n\n")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Versão 1.0  •  Junho 2026")
run.font.size  = Pt(12)
run.font.color.rgb = CINZA_TEXTO

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Engenharia de Teste Industrial — Salcomp Manaus")
run.font.size  = Pt(11)
run.font.color.rgb = CINZA_TEXTO

doc.add_page_break()

# ── 1. VISÃO GERAL ────────────────────────────────────────────────────────────
add_heading(doc, "1. Visão Geral do Produto", 1)
add_body(doc,
    "O MES Client é uma aplicação Windows projetada para operação contínua (24h/7d) em "
    "estações de teste industrial. Sua função é coletar automaticamente os resultados de "
    "teste gerados pelos equipamentos PCM Tester e CYG/NAVAJO, processar esses dados e "
    "armazená-los em um banco de dados PostgreSQL centralizado para rastreabilidade e análise de qualidade.")

add_heading(doc, "1.1 Funcionalidades Principais", 2)
bullets = [
    "Monitoramento automático de pastas de log com detecção de novos arquivos CSV",
    "Parser inteligente com detecção de formato (PCM Tester / CYG)",
    "Insert/Upsert no PostgreSQL com deduplicação por arquivo e linha de origem",
    "Sincronização de arquivos para servidor de rede (modo diff, copy ou sync)",
    "Validação de limites de especificação contra arquivo de referência (spec_limits.csv)",
    "Fila offline com reprocessamento automático em caso de falha de conexão",
    "Interface de sistema de bandeija (system tray) com janelas de Status e Configuração",
    "Controle de acesso por perfil: OPERADOR (visualização) e ENGENHARIA (edição completa)",
    "Criação automática do banco e tabelas na primeira execução",
]
for b in bullets:
    add_bullet(doc, b)

add_heading(doc, "1.2 Arquitetura do Sistema", 2)
add_body(doc,
    "O sistema é composto por módulos independentes que se comunicam através de objetos de "
    "estado thread-safe, permitindo operação robusta em ambiente multithread:")

add_table(doc,
    ["Módulo", "Arquivo", "Responsabilidade"],
    [
        ["Monitor",      "monitor/file_monitor.py",  "Loop principal de varredura de pastas e orquestração"],
        ["Parser",       "parser/cyg_parser.py",     "Leitura e interpretação dos arquivos CSV de teste"],
        ["DB Writer",    "database/db_writer.py",    "Conexão e escrita no PostgreSQL"],
        ["Sync",         "sync/file_sync.py",        "Sincronização de arquivos para destino de rede"],
        ["Spec Check",   "spec/spec_validator.py",   "Validação de limites LSL/USL vs. especificação"],
        ["Buffer",       "buffer/queue_buffer.py",   "Fila offline em JSONL para falhas de conexão"],
        ["UI",           "system/ui_main.py",        "Interface system tray, login, status e configuração"],
        ["Config",       "config/loader.py",         "Carregamento e salvamento de config.yaml"],
        ["State",        "state/app_context.py",     "Estado runtime thread-safe compartilhado"],
    ],
    col_widths=[3, 5, 8]
)

doc.add_page_break()

# ── 2. REQUISITOS ─────────────────────────────────────────────────────────────
add_heading(doc, "2. Requisitos do Sistema", 1)

add_heading(doc, "2.1 Hardware Mínimo", 2)
add_table(doc,
    ["Componente", "Mínimo", "Recomendado"],
    [
        ["Processador", "Intel Core i3 ou equivalente", "Intel Core i5 ou superior"],
        ["Memória RAM",  "4 GB",                        "8 GB"],
        ["Armazenamento","10 GB livres",                "SSD 50 GB+"],
        ["Rede",         "100 Mbps",                   "Gigabit Ethernet"],
    ],
    col_widths=[4, 5, 6]
)

add_heading(doc, "2.2 Software", 2)
add_table(doc,
    ["Software", "Versão", "Observação"],
    [
        ["Windows",    "10 / 11 (64-bit)",   "Obrigatório"],
        ["Python",     "3.10 ou superior",   "Incluído no pacote de instalação"],
        ["PostgreSQL", "13 ou superior",     "Pode ser local ou servidor remoto"],
        ["psycopg2",   "2.9+",               "Driver Python para PostgreSQL (incluído)"],
        ["Pillow",     "10+",                "Biblioteca de imagem (incluída)"],
        ["pystray",    "0.19+",              "Ícone na bandeja do sistema (incluído)"],
        ["PyYAML",     "6+",                 "Leitura de configuração (incluída)"],
    ],
    col_widths=[4, 4, 7]
)

doc.add_page_break()

# ── 3. INSTALAÇÃO ─────────────────────────────────────────────────────────────
add_heading(doc, "3. Instalação e Configuração Inicial", 1)

add_heading(doc, "3.1 Preparação do PostgreSQL", 2)
add_body(doc,
    "Antes de executar o MES Client pela primeira vez, é necessário criar o usuário do banco "
    "de dados. O banco e as tabelas são criados automaticamente pelo aplicativo na primeira execução.")
add_body(doc, "Execute o seguinte comando no psql como superusuário (postgres):", bold=True)
add_code_block(doc, "CREATE USER mes_user WITH PASSWORD 'sua_senha_aqui' CREATEDB;")
add_body(doc,
    "O usuário precisa da permissão CREATEDB apenas para a primeira execução. Após o banco "
    "ser criado, essa permissão pode ser removida por segurança.")

add_heading(doc, "3.2 Configuração do Arquivo .env", 2)
add_body(doc,
    "Crie um arquivo chamado .env na raiz do aplicativo (mesma pasta do config.yaml) "
    "com o seguinte conteúdo:")
add_code_block(doc, "MES_DB_PASSWORD=sua_senha_aqui")
add_body(doc,
    "Este arquivo NUNCA deve ser compartilhado ou enviado para repositórios. "
    "A senha do banco é referenciada no config.yaml como ${MES_DB_PASSWORD} e resolvida "
    "em tempo de execução a partir do .env.", italic=True)

add_heading(doc, "3.3 Configuração do config.yaml", 2)
add_body(doc, "Edite o arquivo config.yaml ajustando os parâmetros para sua estação:")
add_code_block(doc,
"""database:
  enabled: true
  host: localhost          # IP ou hostname do servidor PostgreSQL
  port: 5432
  name: mes_db
  user: mes_user
  password: ${MES_DB_PASSWORD}
  table: mes_test_results  # Nome da tabela (configurável)

station:
  id: PCM_A06_BR-PCMTEST-03   # Identificador único da estação
  type: PCM_TESTER
  model: A06
  line: NAVAJO

log:
  folder: D:\\Data_Info\\A06   # Pasta onde os CSVs de teste são gerados
  recursive: true

sync:
  enabled: true
  destination_folder: E:\\servidor_salcomp
  mode: diff               # diff | copy | sync

parser:
  scan_interval: 5         # Intervalo de varredura em segundos

spec_check:
  enabled: true
  file: spec_limits.csv

auth:
  operador_password: ""    # Vazio = sem senha para OPERADOR
  engenharia_password: "admin"   # Altere para sua senha""")

add_heading(doc, "3.4 Primeira Execução", 2)
add_body(doc, "Na primeira execução, o sistema realiza automaticamente:")
bullets2 = [
    "Carregamento das variáveis de ambiente do arquivo .env",
    "Verificação/criação do banco de dados configurado",
    "Criação de todas as tabelas e índices necessários",
    "Início do monitoramento da pasta de logs",
    "Exibição da tela de login",
]
for b in bullets2:
    add_bullet(doc, b)

doc.add_page_break()

# ── 4. MANUAL DO USUÁRIO ──────────────────────────────────────────────────────
add_heading(doc, "4. Manual do Usuário", 1)

add_heading(doc, "4.1 Perfis de Acesso", 2)
add_body(doc,
    "O MES Client possui dois perfis de acesso para garantir a integridade da configuração "
    "em ambiente de produção:")
add_table(doc,
    ["Ação", "OPERADOR", "ENGENHARIA"],
    [
        ["Ver STATUS",                "✓", "✓"],
        ["Abrir CONFIGURAÇÃO",        "✓ (somente leitura)", "✓"],
        ["Editar CONFIGURAÇÃO",       "✗ (requer senha)", "✓"],
        ["Abrir LIMITES",             "✓ (somente leitura)", "✓"],
        ["Editar LIMITES",            "✗ (requer senha)", "✓"],
        ["START (iniciar monitor)",   "✓", "✓"],
        ["STOP (parar monitor)",      "✗ (requer senha)", "✓"],
        ["EXIT (fechar aplicativo)",  "✗ (requer senha)", "✓"],
        ["ABOUT",                     "✓", "✓"],
    ],
    col_widths=[7, 4.5, 4.5]
)
add_body(doc,
    "Quando um OPERADOR tenta executar uma ação restrita, o sistema exibe um mini-diálogo "
    "solicitando a senha de ENGENHARIA. Se a senha for correta, a ação é executada sem alterar "
    "o perfil da sessão atual.", italic=True)

add_heading(doc, "4.2 Tela de Login", 2)
add_body(doc,
    "Ao iniciar o aplicativo, a tela de login é exibida. Selecione o perfil desejado e informe "
    "a senha correspondente (se configurada). O perfil OPERADOR, por padrão, não requer senha.")

add_heading(doc, "4.3 Menu da Bandeja do Sistema", 2)
add_body(doc,
    "Após o login, o ícone do MES Client aparece na bandeja do sistema (system tray). "
    "Clique com o botão direito para acessar o menu:")
add_table(doc,
    ["Item de Menu", "Função"],
    [
        ["START",        "Inicia o monitoramento de arquivos e envio ao banco"],
        ["STOP",         "Para o monitoramento (requer ENGENHARIA ou senha)"],
        ["STATUS",       "Abre a tela de status em tempo real"],
        ["CONFIGURAÇÃO", "Abre a tela de configuração do sistema"],
        ["LIMITES",      "Abre o editor de limites de especificação"],
        ["ABOUT",        "Informações sobre a versão do software"],
        ["EXIT",         "Encerra o aplicativo (requer ENGENHARIA ou senha)"],
    ],
    col_widths=[4, 11]
)

add_heading(doc, "4.4 Tela de STATUS", 2)
add_body(doc,
    "Exibe em tempo real os indicadores de operação da estação. Atualização automática a cada "
    "1 segundo. Inclui botões para abrir o arquivo de log atual e a pasta de logs.")

add_heading(doc, "4.5 Tela de CONFIGURAÇÃO", 2)
add_body(doc,
    "Exibe todos os parâmetros do config.yaml em modo de somente leitura. "
    "Para editar, clique em EDITAR (requer senha de ENGENHARIA se perfil for OPERADOR). "
    "Após editar, clique em SALVAR — as alterações entram em vigor imediatamente com "
    "reinício automático do monitor.")

add_heading(doc, "4.6 Tela de LIMITES", 2)
add_body(doc,
    "Editor tabular do arquivo spec_limits.csv. Permite visualizar, adicionar, editar e "
    "deletar linhas de especificação. Cada linha define os limites LSL/USL para um passo "
    "de teste de um modelo específico.")

doc.add_page_break()

# ── 5. REFERÊNCIA TÉCNICA ─────────────────────────────────────────────────────
add_heading(doc, "5. Referência Técnica", 1)

add_heading(doc, "5.1 Estrutura do Banco de Dados", 2)

add_heading(doc, "Tabela: mes_test_results", 3)
add_table(doc,
    ["Coluna", "Tipo", "Descrição"],
    [
        ["id",              "BIGSERIAL PK",  "Identificador único autoincremental"],
        ["created_at",      "TIMESTAMP",     "Data/hora de inserção no banco"],
        ["station_id",      "TEXT",          "Identificador da estação de teste"],
        ["model_name",      "TEXT",          "Modelo do produto testado"],
        ["version_name",    "TEXT",          "Versão do firmware/produto"],
        ["serial_number",   "TEXT",          "Número de série da unidade testada"],
        ["result_status",   "TEXT",          "Resultado: PASS / FAIL"],
        ["test_start_time", "TEXT",          "Hora de início do teste"],
        ["test_stop_time",  "TEXT",          "Hora de fim do teste"],
        ["source_file",     "TEXT",          "Caminho do arquivo CSV de origem"],
        ["source_line_no",  "INTEGER",       "Linha do CSV de origem (deduplicação)"],
        ["schema_hash",     "TEXT",          "Hash do esquema de colunas do CSV"],
        ["row_data",        "JSONB",         "Dados completos da linha em formato JSON"],
    ],
    col_widths=[4, 3.5, 7.5]
)

add_heading(doc, "Tabela: mes_csv_schemas", 3)
add_table(doc,
    ["Coluna", "Tipo", "Descrição"],
    [
        ["schema_hash",         "TEXT UNIQUE", "Hash SHA256 do cabeçalho do CSV"],
        ["model_name",          "TEXT",        "Modelo associado ao esquema"],
        ["version_name",        "TEXT",        "Versão associada"],
        ["source_file_pattern", "TEXT",        "Padrão de nome de arquivo"],
        ["columns_json",        "JSONB",       "Lista de colunas do CSV"],
        ["upper_limits_json",   "JSONB",       "Limites superiores (USL) por coluna"],
        ["lower_limits_json",   "JSONB",       "Limites inferiores (LSL) por coluna"],
        ["units_json",          "JSONB",       "Unidades por coluna"],
        ["first_seen",          "TIMESTAMP",   "Primeira vez que este esquema foi visto"],
        ["last_seen",           "TIMESTAMP",   "Última vez atualizado"],
    ],
    col_widths=[4.5, 3.5, 7]
)

add_heading(doc, "Tabela: mes_spec_mismatches", 3)
add_body(doc,
    "Registra divergências entre os limites configurados no testador e os limites "
    "definidos no arquivo spec_limits.csv.")

add_heading(doc, "5.2 Formato do spec_limits.csv", 2)
add_body(doc, "Arquivo CSV com os seguintes campos:")
add_code_block(doc, "enabled,model,step_key,step_name,unit,lsl,usl")
add_table(doc,
    ["Campo", "Tipo", "Descrição"],
    [
        ["enabled",   "0 ou 1",  "1 = linha ativa, 0 = linha ignorada"],
        ["model",     "Texto",   "Modelo do produto (ex: A06). Vazio = aplica a todos"],
        ["step_key",  "Texto",   "Chave do passo de teste no CSV de log"],
        ["step_name", "Texto",   "Nome legível do passo (para relatórios)"],
        ["unit",      "Texto",   "Unidade de medida (V, A, Ω, μA, etc.)"],
        ["lsl",       "Decimal", "Limite inferior de especificação (Lower Spec Limit)"],
        ["usl",       "Decimal", "Limite superior de especificação (Upper Spec Limit)"],
    ],
    col_widths=[3, 3, 9]
)

add_heading(doc, "5.3 Fila Offline", 2)
add_body(doc,
    "Quando o banco de dados está inacessível, os dados são armazenados localmente no arquivo "
    "offline_queue.jsonl (um registro JSON por linha). O monitor tenta reprocessar a fila "
    "automaticamente em cada ciclo de varredura, assim que a conexão for restabelecida.")

add_heading(doc, "5.4 Arquivo de Offsets", 2)
add_body(doc,
    "O arquivo offsets.json registra a posição de leitura de cada arquivo CSV monitorado. "
    "Isso garante que, em caso de reinício do aplicativo, apenas as linhas novas sejam "
    "processadas, sem duplicação de dados.")

doc.add_page_break()

# ── 6. OPERAÇÃO ───────────────────────────────────────────────────────────────
add_heading(doc, "6. Operação e Manutenção", 1)

add_heading(doc, "6.1 Inicialização Automática com o Windows", 2)
add_body(doc,
    "O instalador automático registra o MES Client no Windows Task Scheduler, garantindo "
    "que o aplicativo inicie automaticamente a cada login, com reinicialização automática "
    "em caso de falha:")
add_code_block(doc,
    "Nome da tarefa: MES_Client_Autostart\n"
    "Gatilho: AtLogOn (a cada login de qualquer usuário)\n"
    "Nível: HighestAvailable (privilégios elevados)\n"
    "Reiniciar em falha: 3 tentativas, intervalo de 1 minuto")
add_body(doc, "Para gerenciar a tarefa manualmente:")
add_code_block(doc,
    "# Ver a tarefa\n"
    "schtasks /Query /TN MES_Client_Autostart\n\n"
    "# Desativar temporariamente\n"
    "schtasks /Change /TN MES_Client_Autostart /DISABLE\n\n"
    "# Reativar\n"
    "schtasks /Change /TN MES_Client_Autostart /ENABLE")

add_heading(doc, "6.2 Logs do Sistema", 2)
add_body(doc,
    "Os logs de operação são gravados na pasta configurada em log.folder do config.yaml. "
    "O nível de log padrão é INFO. Eventos importantes registrados:")
add_table(doc,
    ["Evento", "Nível", "Significado"],
    [
        ["MONITOR INICIADO",                    "INFO",    "Monitor de arquivos ativo"],
        ["Conexão com banco estabelecida",       "INFO",    "PostgreSQL acessível"],
        ["Lote inserido com sucesso",            "INFO",    "Dados enviados ao banco"],
        ["Falha ao conectar no banco",           "ERROR",   "PostgreSQL inacessível — fila offline ativada"],
        ["Rollback executado no banco",          "WARNING", "Transação revertida por erro"],
        ["Não foi possível validar/criar banco", "WARNING", "Problema de permissão no PostgreSQL"],
    ],
    col_widths=[6.5, 2.5, 6]
)

add_heading(doc, "6.3 Indicador de Status no Tray", 2)
add_table(doc,
    ["Cor do Ícone", "Significado"],
    [
        ["🟡 Amarelo", "Monitor iniciado, aguardando ou operação normal"],
        ["🟢 Verde",   "Dados processados e enviados ao banco com sucesso"],
        ["🔴 Vermelho", "Erro de conexão com banco ou falha crítica"],
    ],
    col_widths=[5, 10]
)

add_heading(doc, "6.4 Procedimento de Alteração de Configuração", 2)
add_body(doc, "Para alterar parâmetros em produção:", bold=True)
bullets3 = [
    "Abrir o menu de bandeja → CONFIGURAÇÃO",
    "Clicar em EDITAR (informar senha de ENGENHARIA se necessário)",
    "Realizar as alterações nos campos desejados",
    "Clicar em SALVAR — o monitor é reiniciado automaticamente com as novas configurações",
    "Verificar na tela STATUS que o monitor voltou ao estado RUNNING",
]
for b in bullets3:
    add_bullet(doc, b)

add_heading(doc, "6.5 Instalador Automático (Inno Setup)", 2)
add_body(doc,
    "O MES Client inclui um instalador profissional gerado com Inno Setup 6, que automatiza "
    "todo o processo de implantação em novas estações PCM Tester:")
add_table(doc,
    ["Etapa", "O que o instalador faz automaticamente"],
    [
        ["Wizard — Pág. 1", "Coleta: Modelo (A06, A17...) e ID da máquina (BR-PCMTEST-01...)"],
        ["Wizard — Pág. 2", "Coleta: Linha de produção e pasta dos CSVs do TestPad"],
        ["Wizard — Pág. 3", "Coleta: host, porta, banco, usuário e senha do PostgreSQL"],
        ["Wizard — Pág. 4", "Exibe resumo completo da configuração para confirmação"],
        ["Instalação",      "Copia MES_Client.exe e spec_limits.csv para C:\\Utility\\MES"],
        ["Pós-instalação",  "Gera config.yaml personalizado para a estação"],
        ["Pós-instalação",  "Cria tarefa no Task Scheduler (auto-start no login)"],
        ["Pós-instalação",  "Cria atalho na Área de Trabalho (todos os usuários)"],
        ["Conclusão",       "Opção de iniciar o MES Client imediatamente"],
    ],
    col_widths=[4, 11]
)
add_body(doc,
    "O arquivo de distribuição é um único executável de ≈22 MB: MES_Client_Setup_v1.0.exe. "
    "O instalador também inclui desinstalador integrado acessível pelo Painel de Controle → "
    "Programas e Recursos.")

doc.add_page_break()

# ── 7. SEGURANÇA ──────────────────────────────────────────────────────────────
add_heading(doc, "7. Segurança e Boas Práticas", 1)

add_heading(doc, "7.1 Gerenciamento de Senhas", 2)
add_table(doc,
    ["Senha", "Localização", "Observação"],
    [
        ["Banco de dados (PostgreSQL)", "Arquivo .env  →  MES_DB_PASSWORD", "Nunca inserir em config.yaml em texto plano"],
        ["Perfil OPERADOR",             "config.yaml  →  auth.operador_password", "Pode ser vazio (sem senha)"],
        ["Perfil ENGENHARIA",           "config.yaml  →  auth.engenharia_password", "Alterar antes do deploy em produção"],
    ],
    col_widths=[5, 5.5, 4.5]
)

add_heading(doc, "7.2 Recomendações de Segurança", 2)
rec = [
    "Alterar a senha padrão de ENGENHARIA antes de colocar em produção",
    "O arquivo .env não deve ser incluído em backups compartilhados ou repositórios",
    "Restringir o acesso ao diretório de instalação ao usuário de serviço do Windows",
    "O usuário mes_user do PostgreSQL deve ter apenas as permissões necessárias (SELECT, INSERT, UPDATE, CREATE TABLE no banco mes_db)",
    "Realizar backup periódico do banco de dados PostgreSQL",
    "Monitorar os logs para detectar falhas de conexão recorrentes",
]
for r in rec:
    add_bullet(doc, r)

doc.add_page_break()

# ── 8. SUPORTE ────────────────────────────────────────────────────────────────
add_heading(doc, "8. Suporte Técnico", 1)
add_body(doc,
    "Para suporte técnico, dúvidas sobre configuração ou solicitação de novas funcionalidades, "
    "entre em contato com a equipe de Engenharia de Teste:")

add_table(doc,
    ["Contato", "Informação"],
    [
        ["Responsável",  "Roberto Parente"],
        ["Setor",        "Engenharia de Teste — Salcomp Manaus"],
        ["E-mail",       "robertotec.eng3@gmail.com"],
        ["Versão do Software", "MES Client v1.0 — Junho 2026"],
    ],
    col_widths=[5, 10]
)

add_body(doc, "\n")
add_body(doc,
    "Este documento é de uso interno e confidencial. Reprodução ou distribuição não autorizada "
    "é proibida.", italic=True)

doc.add_page_break()

# ── 9. HISTÓRICO DE VERSÕES ───────────────────────────────────────────────────
add_heading(doc, "9. Histórico de Versões", 1)

add_heading(doc, "v1.0.0 — Junho 2026 (Release Inicial de Produção)", 2)
add_body(doc, "Novidades e correções desta versão:", bold=True)
add_table(doc,
    ["Tipo", "Descrição"],
    [
        ["NOVO",     "Instalador profissional Inno Setup com wizard dark premium"],
        ["NOVO",     "Sistema de foco modal completo: focus lock + grab chain + transient"],
        ["NOVO",     "Botão de ajuda (?) na tela LIMITES"],
        ["NOVO",     "Comentários extensivos em todos os módulos Python"],
        ["NOVO",     "Inicialização via Task Scheduler (mais robusto que pasta Startup)"],
        ["NOVO",     "Ícone ICO multi-size gerado programaticamente"],
        ["CORRIGIDO","AJUDA renderizava só a primeira seção (TypeError silencioso no Tkinter)"],
        ["CORRIGIDO","Botão '?' retornava erro 'feche a janela atual'"],
        ["CORRIGIDO","Tela de autenticação piscando ao abrir (conflito de focus lock)"],
        ["CORRIGIDO","EXIT congelava o aplicativo (grab_set em janela oculta)"],
        ["CORRIGIDO","Scripts PowerShell com erros de codificação UTF-8 no PS 5.1"],
    ],
    col_widths=[3, 12]
)

# ── Salvar ────────────────────────────────────────────────────────────────────
output_path = r"D:\MES_Client_Complete\MES_Client_Documentacao_Oficial.docx"
doc.save(output_path)
print(f"Documentação gerada: {output_path}")
